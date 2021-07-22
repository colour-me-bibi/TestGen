#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""TestGen.py: This program does things."""

__author__ = "Benjamin Foreman (bennyforeman1@gmail.com)"


import itertools as it
import random
from collections import defaultdict
from dataclasses import dataclass, field
from enum import Enum, auto
from string import ascii_uppercase

from docx import Document  # pip install python-docx


class QuestionType(Enum):
    BOOLEAN = auto()
    MULTIPLE_CHOICE = auto()
    SELECT_ALL = auto()
    SHORT_ANSWER = auto()

    @staticmethod
    def from_str(label: str):  # TODO could maybe be better
        label = label.replace(" ", "").lower()

        if "bool" in label:
            return QuestionType.BOOLEAN

        if "multiple" in label:
            return QuestionType.MULTIPLE_CHOICE

        if "select" in label:
            return QuestionType.SELECT_ALL

        if "short" in label:
            return QuestionType.SHORT_ANSWER

        raise NotImplementedError


@dataclass
class Answer:
    text: str
    is_correct: bool = False


@dataclass
class Question:
    prompt: str
    question_type: QuestionType
    boolean: bool = None
    answers: list[Answer] = field(default_factory=list)
    response: str = None  # could be EEIs as a list of strings
    is_required: bool = False


@dataclass
class Test:
    title: str
    test_id: int = 0
    questions: list[Question] = field(default_factory=list)


def get_input(prompt: str, cast=lambda x: x, validation=lambda x: True):
    while True:
        try:
            response = cast(input(prompt))

            if not validation(response):
                raise ValueError

            return response
        except Exception:
            print(f"{response} is not a valid input.")


def generate_n_tests(test: Test, n: int):
    grouped_is_required = defaultdict(list)
    for question in test.questions:
        grouped_is_required[question.is_required].append(question)

    required, nonrequired = grouped_is_required[True], grouped_is_required[False]

    grouped_type = defaultdict(list)
    for question in nonrequired:
        grouped_type[question.question_type].append(question)

    boolean, multiple_choice, select_all, short_answer = (
        grouped_type[QuestionType.BOOLEAN],
        grouped_type[QuestionType.MULTIPLE_CHOICE],
        grouped_type[QuestionType.SELECT_ALL],
        grouped_type[QuestionType.SHORT_ANSWER],
    )

    print(f"{len(test.questions)} total questions")
    print(f"\t{len(required)} required questions")
    print(f"\t{len(nonrequired)} nonrequired questions")
    print(f"\t\t{len(boolean)} boolean questions")
    print(f"\t\t{len(multiple_choice)} multiple_choice questions")
    print(f"\t\t{len(select_all)} select_all questions")
    print(f"\t\t{len(short_answer)} short_answer questions")

    if len(boolean):
        boolean_count = get_input(
            f"Require how many boolean questions (max {len(boolean)}) -> ", int, lambda x: x <= len(boolean)
        )
    if len(multiple_choice):
        multiple_choice_count = get_input(
            f"Require how many multiple_choice questions (max {len(multiple_choice)}) -> ",
            int,
            lambda x: x <= len(multiple_choice),
        )
    if len(select_all):
        select_all_count = get_input(
            f"Require how many select_all questions (max {len(select_all)}) -> ", int, lambda x: x <= len(select_all)
        )
    if len(short_answer):
        short_answer_count = get_input(
            f"Require how many short_answer questions (max {len(short_answer)}) -> ",
            int,
            lambda x: x <= len(short_answer),
        )

    tests = list()
    for i in range(1, n + 1):
        questions = (
            required
            + random.sample(boolean, int(boolean_count))
            + random.sample(multiple_choice, int(multiple_choice_count))
            + random.sample(select_all, int(select_all_count))
            + random.sample(short_answer, int(short_answer_count))
        )

        for question in questions:
            random.shuffle(question.answers)
        random.shuffle(questions)

        tests.append(Test(title=test.title, questions=questions, test_id=i))

    return tests


def tests_to_doc(tests: list[Test]):
    title = tests[0].title

    document = Document()

    for test_idx, test in enumerate(tests, 1):
        document.add_heading(f"{title} (ID: {test_idx})", 0)

        for question_idx, question in enumerate(test.questions, 1):
            document.add_heading(f"{question_idx} - {question.prompt}", level=1)

            if question.question_type == QuestionType.BOOLEAN:
                document.add_paragraph("\tTrue or False")
            elif question.question_type == QuestionType.MULTIPLE_CHOICE:
                document.add_paragraph("Multiple choice:")

                for answer_idx, answer in enumerate(question.answers):
                    document.add_paragraph(f"\t{ascii_uppercase[answer_idx]} - {answer.text}")
            elif question.question_type == QuestionType.SELECT_ALL:
                document.add_paragraph("Select all that apply:")

                for answer_idx, answer in enumerate(question.answers):
                    document.add_paragraph(f"\t{ascii_uppercase[answer_idx]} - {answer.text}")
            elif question.question_type == QuestionType.SHORT_ANSWER:
                for _ in range(2):
                    document.add_paragraph()
            else:
                raise NotImplementedError

        document.add_page_break()

    for test_idx, test in enumerate(tests, 1):
        document.add_heading(f"{title} (ID: {test_idx})", 0)

        for question_idx, question in enumerate(test.questions, 1):
            answer_text = None
            if question.question_type == QuestionType.BOOLEAN:
                answer_text = str(question.boolean)
            elif question.question_type == QuestionType.MULTIPLE_CHOICE:
                answer_text = " or ".join(ascii_uppercase[i] for i, answer in enumerate(question.answers) if answer.is_correct)
            elif question.question_type == QuestionType.SELECT_ALL:
                answer_text = " and ".join(ascii_uppercase[i] for i, answer in enumerate(question.answers) if answer.is_correct)
            elif question.question_type == QuestionType.SHORT_ANSWER:
                answer_text = question.response
            else:
                raise NotImplementedError

            document.add_paragraph(f"{question_idx} - {answer_text}")

        document.add_paragraph()
        document.add_page_break()

    document.save(f"{title}-{len(tests)}.docx")


def parse_marked_string(string: str):
    is_marked = string.startswith("-")
    text = string.lstrip("-").lstrip() if is_marked else string

    return text, is_marked


def lcount(string: str, *patterns: list[str]):
    return max(len(string) - len(string.lstrip(pattern)) for pattern in patterns)


def parse_txt_to_test(filepath: str):
    with open(filepath) as rf:
        lines = (line for line in (line.rstrip() for line in rf.readlines()) if line)

    groups = (
        (key, [x.lstrip() for x in group])
        for key, group in it.groupby(lines, key=lambda line: lcount(line, " ", "\t"))
    )

    next(groups)  # title marker
    title = title=list(next(groups)[1])[0].lstrip()

    questions = list()

    question_type = None
    while (group := next(groups, None)) is not None:
        if group[0] == 0:
            question_type = QuestionType.from_str(group[1][0])
        else:
            prompt, is_required = parse_marked_string(group[1][0])
            answers = next(groups)[1]

            kwargs = {
                "prompt": prompt,
                "question_type": question_type,
                "is_required": is_required,
            }

            if question_type == QuestionType.BOOLEAN:
                kwargs["boolean"] = "true" in answers[0].lower()
            elif question_type in (QuestionType.MULTIPLE_CHOICE, QuestionType.SELECT_ALL):
                kwargs["answers"] = [Answer(*parse_marked_string(string)) for string in answers]
            elif question_type == QuestionType.SHORT_ANSWER:
                kwargs["response"] = answers[0]
            else:
                raise NotImplementedError

            questions.append(Question(**kwargs))

    return Test(title=title, questions=questions)


if __name__ == "__main__":
    testbank = parse_txt_to_test("example_test.txt")
    tests_to_doc(generate_n_tests(testbank, 2))
